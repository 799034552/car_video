from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from PyQt5.QtWidgets import (QWidget, QPushButton, 
    QHBoxLayout, QVBoxLayout, QApplication,QLabel,QLineEdit, QCheckBox,QTextEdit)
from PyQt5.QtCore import QTimer
from PyQt5 import QtCore
from PyQt5.Qt import (QThread, pyqtSignal)
from time import sleep
import time, datetime
from PIL import Image
from PyQt5.QtGui import QTextCursor
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import xlrd
import xlwt
import xlutils.copy as xlCopy
import os,sys,json,requests
import math
from selenium.webdriver.chrome.options import Options
# 设置默认字体
def chg_font(obj, fontname='微软雅黑', size=None):
    ## 设置字体函数
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
    if size and isinstance(size, Pt):
        obj.font.size = size
# 读取车牌与进度
def readData(carFile, hisFile,isCon):
    if not isCon or not os.path.exists(hisFile):
        lastFile = xlrd.open_workbook(carFile)
        sheet = lastFile.sheet_by_index(0)
        res = []
        for i in range(1,sheet.nrows):
            res.append(sheet.row_values(i))
        return res,0,0
    else:
        lastFile = xlrd.open_workbook(hisFile)
        sheet = lastFile.sheet_by_index(0)
        res = []
        good = 0
        index = -1
        for i in range(1,sheet.nrows):
            if sheet.row_values(i)[1] == "正常":
                good += 1
            elif index == -1 and len(sheet.row_values(i)[1]) == 0:
                index = i
            res.append(sheet.row_values(i))
        return res,index - 1,good


def readFile(file = "车牌.xlsx"):
    lastFile = xlrd.open_workbook(file)
    sheet = lastFile.sheet_by_index(0)
    res = []
    for i in range(0,sheet.nrows):
        res.append(sheet.row_values(i))
    return res

# 导出数据
def writeData(name, corName = "", resName="结果.docx"):
    try:
        doc = Document(resName)
    except:
        doc = Document()
    #chg_font(doc.styles['Normal'], fontname='宋体')
    tab = doc.add_table(rows=7, cols=2, style="Table Grid")  # 添加一个4行4列的空表
    tab.style.paragraph_format.space_before = Pt(0)
    tab.style.paragraph_format.space_after = Pt(0)
    tab.cell(0, 0).merge(tab.cell(0, 1))
    tab.cell(1, 0).merge(tab.cell(1, 1))
    tab.cell(2, 0).paragraphs[0].add_run("查验日期")
    
    tab.cell(2, 1).paragraphs[0].add_run(time.strftime("%Y-%m-%d", time.localtime()) )
    tab.cell(3, 0).paragraphs[0].add_run("所属运输公司")
    tab.cell(3, 1).paragraphs[0].add_run(corName)
    tab.cell(4, 0).paragraphs[0].add_run("车牌号")
    tab.cell(4, 1).paragraphs[0].add_run(name)
    tab.cell(5, 0).paragraphs[0].add_run("视频轨迹情况")
    tab.cell(6, 0).paragraphs[0].add_run("审验人")
    tab.cell(0, 0).paragraphs[0].add_run().add_picture("./img/轨迹.png", width=Inches(5.75))
    tab.cell(1, 0).paragraphs[0].add_run().add_picture("./img/实时视频_处理后.png", width=Inches(5.75))
    for i in range(len(tab.rows)):
        for j in range(len(tab.columns)):
            for item in tab.cell(i, j).paragraphs:
                item.paragraph_format.space_before=Pt(0)
                item.paragraph_format.space_after = Pt(0)
    doc.save(resName)

# 写出结果excel
def writeRes(file, index, car, msg):
    line = [0]
    def writeLine(line, sheet,*arr):
        col = 0
        for i in arr:
            sheet.write(line[0], col, i)
            col += 1
        line[0] += 1
    if os.path.exists(file):
        resFile = xlrd.open_workbook(file)
        resFile = xlCopy.copy(resFile)
        resExcel = resFile.get_sheet(0)
    else:
        resFile = xlwt.Workbook()
        resExcel = resFile.add_sheet('sheet1')
        writeLine(line, resExcel, ["车牌", "状态"])
        for item in car:
            writeLine(line, resExcel, item)
    resExcel.write(index+1, 1, msg)
    resFile.save(file)




def getByName(driver, name = "粤SGH6593"):
    waitNotEle(driver,".ant-spin-dot.ant-spin-dot-spin")
    inputTxt = driver.find_elements_by_css_selector(".ant-input.ant-input-lg")[0]
    inputTxt.clear()
    inputTxt.send_keys(name)
    sleep(1)
    # 等待搜索完毕
    while len(driver.find_elements_by_css_selector(".ant-tree-treenode-checkbox-checked")) > 20:
        sleep(0.1)
        
    # 查找主列表
    waitEle(driver,".ant-tree-treenode-checkbox-checked", 2)
    selectList = driver.find_elements_by_css_selector(".ant-tree-treenode-checkbox-checked")
    if (len(selectList) < 1):
        return {
            "state": "error",
            "msg": "没有找到该车牌的车"
        }
    eleClass = selectList[1].get_attribute("class")
    if (eleClass.find("ant-tree-treenode-switcher-open") == -1):
        selectList[1].click()
    # 获取公司名
    cor = driver.find_elements_by_css_selector(".ant-tree-title")[1].text
    cor = cor[0:cor.find("(")]
    # sleep(2)
    waitEle(driver,".ant-tree-node-content-wrapper.ant-tree-node-content-wrapper-normal")
    # 查找子列表
    selectList = driver.find_elements_by_css_selector(".ant-tree-node-content-wrapper.ant-tree-node-content-wrapper-normal")
    waitEle(driver,".ant-tree-node-content-wrapper.ant-tree-node-content-wrapper-normal")

    pic = driver.find_elements_by_css_selector(".ant-tree-node-content-wrapper.ant-tree-node-content-wrapper-normal img")[0]
    picSrc = pic.get_attribute("src")
    if (picSrc.find("-online") != -1):
        waitClickEle(driver, ".ant-tree-node-content-wrapper.ant-tree-node-content-wrapper-normal")
        selectList[0].click()
    else:
        return {
            "state": "error",
            "msg": "该车辆没有在运行"
        }
    waitClickEle(driver,".bottom-center.amap-info-contentContainer button")
    button = driver.find_elements_by_css_selector(".bottom-center.amap-info-contentContainer button")
    waitClickEle(driver,".bottom-center.amap-info-contentContainer button")
    button[0].click()
    button[2].click()
    openWebs = {}
    # 判断打开的网页
    for item in driver.window_handles:
        driver.switch_to.window(item)
        if str(driver.current_url).find("monitorMap") != -1:
            openWebs["monitorMap"] = item
        elif str(driver.current_url).find("monitorVideo") != -1:
            openWebs["monitorVideo"] = item
        elif str(driver.current_url).find("trackPlayback") != -1:
            openWebs["trackPlayback"] = item

    qtPrint("正在处理轨迹页面")
    # 处理轨迹页面
    driver.switch_to.window(openWebs["trackPlayback"])
    waitEle(driver, ".amap-icon")
    sleep(4)
    driver.save_screenshot("./img/轨迹.png")
    driver.close()

    qtPrint("正在处理实时视频")
    #处理实时视频
    jsCode = 'let timeInt, isTimeOut = false, maxTimes = 60, srcTimes = 6\n\
    function fun() {\n\
    console.log("left  " + maxTimes.toString())\n\
    let isBad = false\n\
    let hasSrc = false\n\
    maxTimes--\n\
    if (maxTimes <= 0) {\n\
        let a = document.createElement("div")\n\
        a.className = "python-get-class"\n\
        a.setAttribute("type", "error")\n\
        document.querySelector("body").appendChild(a)\n\
        console.log("超时")\n\
        clearInterval(timeInt)\n\
        return\n\
    }\n\
    videos = document.querySelectorAll("video")\n\
    videos.forEach(item => {\n\
        if (item.src) {\n\
            hasSrc = true\n\
            console.log(item.readyState)\n\
            if (item.readyState === 0) {\n\
                isBad = true\n\
            }\n\
        }\n\
    })\n\
    if (!hasSrc) {\n\
        srcTimes--\n\
    }\n\
    if (srcTimes <= 0) {\n\
        let a = document.createElement("div")\n\
        a.className = "python-get-class"\n\
        a.setAttribute("type", "noVideo")\n\
        document.querySelector("body").appendChild(a)\n\
        console.log("没有视频")\n\
        clearInterval(timeInt)\n\
        return\n\
    }\n\
    if (!isBad && hasSrc) {\n\
        let a = document.createElement("div")\n\
        a.className = "python-get-class"\n\
        a.setAttribute("type", "good")\n\
        console.log("所有视频加载完成")\n\
        document.querySelector("body").appendChild(a)\n\
        clearInterval(timeInt)\n\
    }\n\
    }\n\
    timeInt = setInterval(fun, 1000)'
    driver.switch_to.window(openWebs["monitorVideo"])
    waitClickEle(driver, ".ant-col.ant-col-4")
    driver.find_elements_by_css_selector(".ant-col.ant-col-4")[2].click()
    driver.execute_script(jsCode)
    waitEle(driver, ".python-get-class",70, 0.8)
    type = driver.find_elements_by_css_selector(".python-get-class")[0].get_attribute("type")
    if (type == "good"):
        sleep(5)
        driver.save_screenshot("./img/实时视频.png")
        picture = Image.open("./img/实时视频.png")
        picture = picture.crop((617, 115, 617 + 1144, 115 + 620))
        picture.save("./img/实时视频_处理后.png")
        pass
    elif type == "noVideo":
        driver.close()
        driver.switch_to.window(openWebs["monitorMap"])
        return {
            "state": "error",
            "msg": "车辆没有视频（离线）"
        }
    else:
        driver.close()
        driver.switch_to.window(openWebs["monitorMap"])
        return {
            "state": "error",
            "msg": "读取视频超时"
        }
    driver.close()
    driver.switch_to.window(openWebs["monitorMap"])
    return {
            "state": "good",
            "msg": "正常",
            "cor": cor
    }


# 等待浏览器加载元素
def waitEle(driver, ele, times=30, between = 0.3):
    try:
        WebDriverWait(driver, times, between).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, ele))
        )
    except:
        qtPrint("等待元素出错" + ele)

# 等待元素消失
def waitNotEle(driver, ele, times=30):
    try:
        WebDriverWait(driver, times, 0.2).until_not(
            EC.presence_of_element_located((By.CSS_SELECTOR, ele))
        )
    except:
        qtPrint("等待元素出错" + ele)

def waitClickEle(driver, ele, times = 30):
    try:
        WebDriverWait(driver, times, 0.2).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, ele))
        )
    except:
        qtPrint("等待元素出错")


# 可拖拽文本框
class MyLineEdit(QLineEdit):
    def __init__(self):
        super(MyLineEdit, self).__init__()
        self.setAcceptDrops(True)

    # 拖拽内容进来的时候触发这个事件
    def dragEnterEvent(self, e):
        print(e)
        if e.mimeData().hasText():
            e.accept()
        else:
            e.ignore()

    # 内容放置（松开鼠标时）时触发这个事件
    def dropEvent(self, e):
      print(e.mimeData().text())
      self.setText(e.mimeData().text()[8:])

cars = 0;index = 0;good = 0;choose = 0;checkBox2 = 0;resExcel = 0;progressLabel = 0; mySignal = 0;driver= 0
# 在qt上显示调试信息
def qtPrint(s):
    global mySignal
    mySignal.emit(s)
    # debugBox.append(s + "\n")
    # print(s)

# 用线程启动浏览器爬虫，防止界面死锁
class Thread(QThread):
    _signal =pyqtSignal(str)
    def __init__(self):
        global mySignal
        super().__init__()
        mySignal = self._signal
        self.now = False
        self.resDocxName = ""
        self.resExcelName = ""

    def run(self):
        global cars,index,good,choose,checkBox2,resExcel,progressLabel,debugBox, driver
        length = len(cars)
        x = 0
        if not choose:
            x = 1980
        if checkBox2.checkState():
            try:
                os.remove("./"+ getDate()+ "结果_马上开始.docx")
                os.remove("./"+ getDate()+ "结果_马上开始.xls")
            except:
                pass
        opt = Options()
        opt.add_argument('--no-sandbox')                # 解决DevToolsActivePort文件不存在的报错
        opt.add_argument('window-size=1920x3000')       # 设置浏览器分辨率
        opt.add_argument('--disable-gpu')               # 谷歌文档提到需要加上这个属性来规避bug
        # opt.add_argument('--hide-scrollbars')           # 隐藏滚动条，应对一些特殊页面
        # opt.add_argument('blink-settings=imagesEnabled=false')      # 不加载图片，提升运行速度
        # opt.add_argument('--headless')

        # 读取文件
        driver = webdriver.Chrome(options=opt)
        driver.set_window_size(1980,1280)
        driver.set_window_position(x, 0)
        # 打开登陆
        driver.get("http://112.94.64.104:8080/login")
        # 登陆
        waitEle(driver,".ant-input")
        inputBox =  driver.find_elements_by_css_selector(".ant-input")
        inputBox[0].send_keys("zhitou")
        inputBox[1].send_keys("123456")
        waitEle(driver,".ant-btn")
        driver.find_elements_by_class_name("ant-btn")[0].click()
        waitEle(driver,".ant-menu-submenu-title")
        # 跳转到监控页面
        driver.execute_script("window.location.href = 'http://112.94.64.104:8080/monitorMap'")
        for i in range(index, len(cars)):
            try:
                item = cars[i]
                qtPrint("正在开始" + str(item[0]) + "," + str(i + 1) + "/" + str(len(cars)))
                res = getByName(driver, item[0])
                if res["state"] == "good":
                    good += 1
                    writeData(item[0], res["cor"], self.resDocxName)
            except:
                res = {}
                res["msg"] = "未知错误"
            finally:
                print(res)
                qtPrint(res["msg"])
                progressLabel.setText("进度：%d/%d，成功处理：%d/%d" % (i+1,length,good,length))
                writeRes(self.resExcelName, i, cars, res["msg"])
                if i % 3  == 0:
                    driver.refresh()
        qtPrint("程序结束")
        driver.quit()
        driver = 0
        self._signal.emit("quit")


# 获取时间
def getDate():
    nowDate = datetime.datetime.today()
    nowDateStr = str(nowDate.month) + '.' + str(nowDate.day)
    return nowDateStr
class Example(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()

    #按钮事件

    def setTime(self):
        wantHour = [10, 15, 23]


    # 设置结果文件的名字
    def setResName(self, fileDirName = "./结果"):
        wantHour = [10, 15, 23]
        nowDate = datetime.datetime.today()
        nowDateStr = str(nowDate.month) + '.' + str(nowDate.day)
        nowHour = nowDate.hour
        fileName = ""
        for i in range(len(wantHour)):
            if nowHour <= wantHour[i]:
                fileName = nowDateStr + "结果" + str(i + 1) + "次"
                break
        self.resExcel.setText(fileDirName + '/' + fileName + ".xls")
        print(fileDirName + '/' + fileName, nowDateStr)
        return fileDirName + '/' + fileName, nowDateStr


    def initUI(self):
      #初始化按钮
      self.originButton = QPushButton("马上开始")
      self.setTimeButton = QPushButton("定时开始")
      #按钮事件
      self.originButton.clicked.connect(self.start)
      expTime = "2021-10-31 00:00:00"
      isnotEpi, nowTime, timeStr = getTIme(expTime)
      self.setWindowTitle('车机入网认证审核系统')
      # 大于某个时间，程序不可用
      if not isnotEpi:
          self.originButton.setEnabled(False)

      #文本框
      label1=QLabel(self)
      label1.setText('车牌文件(拖入)')
      carExcel=MyLineEdit()
      self.carExcel = carExcel
      carExcel.setText("./车牌.xlsx")
      carExcel.setDragEnabled(True)
      label2=QLabel(self)
      label2.setText('结果excel(拖入)')
      resExcel=MyLineEdit()
      self.resExcel = resExcel
      resExcel.setText("./" + getDate() + "结果_马上开始.xls")
      resExcel.setDragEnabled(True)

      #复选框
      self.checkBox1 = QCheckBox("显示浏览器", self)
      self.checkBox2 = QCheckBox("删除历史（删除两个结果文件，使用新进度）", self)
      hboxCheckBox = QHBoxLayout()
      hboxCheckBox.addWidget(self.checkBox1)
      hboxCheckBox.addWidget(self.checkBox2)
      self.checkBox2.setChecked(False)

      # 进度与调试
      self.progressLabel = QLabel(self)
      self.debugBox = QTextEdit(self)
    #   self.debugBox.setEnabled(False)
      self.debugBox.setText("调试信息")
      hboxProgress = QHBoxLayout()
      # hboxProgress.addWidget(self.debugBox,1, QtCore.Qt.AlignBottom)
    #   hboxProgress.addStretch(1)
     # hboxProgress.addWidget(self.progressLabel,0, QtCore.Qt.AlignBottom)
      self.progressLabel.setText("进度：0/0，成功处理：0/0")


      #文本布局
      hbox = QHBoxLayout()
      hbox.addWidget(label1)
      hbox.addWidget(carExcel)
      hbox11 = QHBoxLayout()
      hbox11.addWidget(label2)
      hbox11.addWidget(resExcel)

      #倒计时
      hbox13 = QHBoxLayout()
      self.leftLabel = QLabel(self)
      self.leftLabel.setText("距离下一次自动运行:")
    #   hbox13.addStretch(1)
      hbox13.addWidget(self.leftLabel)

      #按钮布局
      hbox1 = QHBoxLayout()
      expLable = QLabel(self)
      expLable.setText('data expired：' + expTime)
      expLable.setStyleSheet('QLabel{color :gray}')
      hbox1.addWidget(expLable)
      hbox1.addStretch(1)
      hbox1.addWidget(self.progressLabel)
      hbox1.addWidget(self.originButton)


      # 大标题
      titleLable = QLabel("车机入网认证审核系统")
      titleLable.setStyleSheet("QLabel{font-size:40px;font-weight:bold;margin-bottom:5px;color:red}")

      #整体布局
      vbox = QVBoxLayout()
      vbox.addWidget(titleLable,0,QtCore.Qt.AlignHCenter)
      vbox.addLayout(hbox)
      vbox.addLayout(hbox11)
      vbox.addLayout(hboxCheckBox)
      vbox.addWidget(self.debugBox,1)
      
    #   vbox.addStretch(1)
      vbox.addLayout(hbox13)
      vbox.addLayout(hbox1)
      #启用布局
      self.setLayout(vbox)
      desktop = QApplication.desktop()
      self.move(desktop.width() / 4, desktop.height() *0.1)
      self.resize(desktop.width() / 2, desktop.height() *0.8)

      self.timer = QTimer(self)
      self.timer.start(1000)
      self.leafTime = -1
      self.timer.timeout.connect(self.checkTime)
      self.checkTime()

      #设置窗口信息
      self.show()
    
    def checkTime(self):
        wantHour = [8, 15, 23]
        if (self.leafTime == -1):
            nowStamp = time.time()
            self.wantTime = []
            self.times = -1
            for i in range(len(wantHour)):
                tempTime = time.strftime("%Y-%m-%d "+str(wantHour[i])+":00:00", time.localtime(nowStamp))
                
                tempTimeStamp = time.mktime(time.strptime(tempTime, '%Y-%m-%d %H:%M:%S'))
                self.wantTime.append(tempTimeStamp)
                if self.leafTime == -1 and nowStamp < tempTimeStamp:
                    self.times = i + 1
                    self.leafTime = math.ceil((tempTimeStamp - nowStamp) + 1)
        self.leafTime -= 1
        if self.leafTime == 0:
            self.smallStart()
            nowStamp = time.time()
            for item in self.wantTime:
                if nowStamp < item:
                    self.leafTime = math.ceil((item - nowStamp) + 1)
                    break
        if self.leafTime != -2:
            self.leftLabel.setText("距离下一次自动运行："+ str(datetime.timedelta(seconds=self.leafTime)))
        else:
            self.leafTime = -1
            self.leftLabel.setText("无需自动运行")
            self.timer.stop()

    def getTimes(self):
        nowStamp = time.time()
        times = 0
        for i in range(len(self.wantTime)):
            if nowStamp < self.wantTime[i] + 1000:
                times = i + 1
        return times
                
    def smallStart(self):
        if self.originButton.isEnabled == False:
            return
        global cars,index,good,choose,checkBox2,resExcel,progressLabel,debugBox
        if (self.times != -1):
            cars,index,good = readData(self.carExcel.text(), "./" + getDate() + "结果_"+str(self.getTimes())+"次.xls", not self.checkBox2.checkState())
            length = len(cars)
            if index < 0:
                index = len(cars)
            self.progressLabel.setText("进度：%d/%d，成功处理：%d/%d" % (index,length,good,length))

            # 处理用户选择
            choose = self.checkBox1.checkState()
            checkBox2 = self.checkBox2
            resExcel = self.resExcel
            progressLabel = self.progressLabel
            debugBox = self.debugBox
            
            self.thread = Thread()
            self.thread._signal.connect(self.ThreadSignal)
            self.thread.resDocxName = "./" + getDate() + "结果_"+str(self.getTimes())+"次.docx"
            self.thread.resExcelName = "./" + getDate() + "结果_"+str(self.getTimes())+"次.xls"
            self.thread.start()
            self.originButton.setEnabled(False)
    
      
    def start(self):

        global cars,index,good,choose,checkBox2,resExcel,progressLabel,debugBox
        # 读取历史文件
        cars,index,good = readData(self.carExcel.text(), "./" + getDate() + "结果_马上开始.xls", not self.checkBox2.checkState())
        length = len(cars)
        if index < 0:
            index = len(cars)
        self.progressLabel.setText("进度：%d/%d，成功处理：%d/%d" % (index,length,good,length))

        # 处理用户选择
        choose = self.checkBox1.checkState()
        checkBox2 = self.checkBox2
        resExcel = self.resExcel
        progressLabel = self.progressLabel
        debugBox = self.debugBox
        
        self.thread = Thread()
        self.thread._signal.connect(self.ThreadSignal)
        self.thread.resDocxName = "./" + getDate() + "结果_马上开始.docx"
        self.thread.resExcelName = "./" + getDate() + "结果_马上开始.xls"
        self.thread.start()
        self.originButton.setEnabled(False)
        pass

    # 显示调试信息
    def qtPrint(self, s):
        self.debugBox.append(s)
        print(s)
        self.debugBox.moveCursor(QTextCursor.End)

    # 调试信号接收
    def ThreadSignal(self, s):
        if s != "quit":
            self.qtPrint(s)
        else:
            self.originButton.setEnabled(True)
    # 退出事件，关闭浏览器
    def closeEvent(self, event):
        global driver
        if driver != 0:
            driver.quit()
        event.accept()

def getTIme(timeStr = "2021-9-30 00:00:00"):
    try:
        tabaoTime = requests.get("http://api.m.taobao.com/rest/api3.do?api=mtop.common.getTimestamp")
        tabaoTime = json.loads(tabaoTime.text)
        nowTimeSt = int(tabaoTime["data"]["t"])/ 1000
        nowTime = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(nowTimeSt))
        expTimeStamp=time.mktime(time.strptime(timeStr, '%Y-%m-%d %H:%M:%S'))
        return (expTimeStamp > nowTimeSt), nowTime, timeStr
    except:
        return False, "错误", "错误"
if __name__ == '__main__':
    
    app = QApplication(sys.argv)
    ex = Example()
    sys.exit(app.exec_())




    